"""
Generate grammar_teaching.docx from the 15 grammar units.
Run:  python make_word.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

# ── Colours ────────────────────────────────────────────────────────────────
PURPLE   = RGBColor(0x6d, 0x28, 0xd9)
TEAL     = RGBColor(0x0d, 0x94, 0x88)
AMBER    = RGBColor(0xb4, 0x53, 0x09)
DARK_PUR = RGBColor(0x1e, 0x1b, 0x4b)
GRAY     = RGBColor(0x6b, 0x72, 0x80)
BLACK    = RGBColor(0x1f, 0x29, 0x37)
RED      = RGBColor(0xb9, 0x1c, 0x1c)

BG_PUR   = RGBColor(0xed, 0xe9, 0xfe)  # light purple bg
BG_TEAL  = RGBColor(0xd1, 0xfa, 0xe5)
BG_AMB   = RGBColor(0xfe, 0xf3, 0xc7)  # amber light


TIER_COLOR = {1: PURPLE, 2: TEAL, 3: AMBER}

# ── Helper: set paragraph shading ──────────────────────────────────────────
def set_shading(paragraph, fill_hex: str):
    """Add background shading to a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    pPr.append(shd)

def set_cell_shading(cell, fill_hex: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def add_border_bottom(paragraph, color_hex='C4B5FD', size=4):
    """Add a bottom border to a paragraph (used for section separators)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), str(size))
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color_hex)
    pBdr.append(bot)
    pPr.append(pBdr)

def set_all_borders(table, color='C4B5FD', sz=4):
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for side in ('top','left','bottom','right','insideH','insideV'):
                el = OxmlElement(f'w:{side}')
                el.set(qn('w:val'), 'single')
                el.set(qn('w:sz'), str(sz))
                el.set(qn('w:space'), '0')
                el.set(qn('w:color'), color)
                tcBorders.append(el)
            tcPr.append(tcBorders)

# ── Helper: add run with formatting ────────────────────────────────────────
def add_run(para, text, bold=False, italic=False, color=None, size=None,
            underline=False, highlight=False):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    return run

# ── Parse inline HTML with <em> tags ───────────────────────────────────────
def add_html_text(para, html, base_color=None, base_size=None, em_color=None):
    """Render text containing <em>…</em> highlights into a paragraph."""
    if em_color is None:
        em_color = PURPLE
    parts = re.split(r'<em>(.*?)</em>', html, flags=re.DOTALL)
    for i, part in enumerate(parts):
        part = part.replace('&amp;', '&').replace('&nbsp;', ' ').replace('<br>', '\n')
        part = re.sub(r'<[^>]+>', '', part)  # strip remaining tags
        if i % 2 == 0:
            if part:
                r = para.add_run(part)
                if base_color:
                    r.font.color.rgb = base_color
                if base_size:
                    r.font.size = Pt(base_size)
        else:
            if part:
                r = para.add_run(part)
                r.bold = True
                r.underline = True
                r.font.color.rgb = em_color
                if base_size:
                    r.font.size = Pt(base_size)

# ── Data ───────────────────────────────────────────────────────────────────
UNITS = [
  {
    "num": 1, "tier": 1, "freq": "12/13",
    "zh": "現在完成式",
    "en": "Present Perfect Tense",
    "formula": "S + has / have + 過去分詞（p.p.）",
    "struct_note": "常見搭配：since（自從）、for（達…之久）、already（已經）、yet（尚未）、ever（曾經）、never（從未）",
    "rules": [
        "第三人稱單數主詞用 has，其餘主詞用 have",
        "表示過去動作對現在仍有影響，或從過去持續到現在的狀態",
        "現在完成進行式：has / have + been + V-ing，強調動作持續進行中",
    ],
    "examples": [
        ("Ms. Johnson <em>has been taking</em> phone calls since she entered the office this morning.",
         "強森女士從今天早上進辦公室起，就一直在接電話。"),
        ("I <em>have tried</em> to call you many times, but you <em>have never answered</em>.",
         "我已經打電話給你很多次了，但你從沒接聽。"),
        ("For the past twenty years, my father <em>has worked</em> in a school library.",
         "過去二十年來，我父親一直在學校圖書館工作。"),
    ],
    "q_stem": "Mary and her friends _____ at the restaurant for two hours. They are very hungry now.",
    "options": ["(A) wait", "(B) waited", "(C) have been waiting", "(D) are waiting"],
    "answer": "(C) have been waiting",
    "explanation": "for two hours 說明動作持續了一段時間，且 They are very hungry now 表示動作的影響延續至現在，應使用現在完成進行式 have been waiting。(A) 原形、(B) 過去式、(D) 現在進行式均不符合「持續到現在」的語意。",
  },
  {
    "num": 2, "tier": 1, "freq": "11/13",
    "zh": "間接問句／引述語氣",
    "en": "Reported Speech / Indirect Questions",
    "formula": "疑問詞 / if / whether + 主詞 + 動詞（直述語序）",
    "struct_note": "重點：間接問句語序改回「主詞 + 動詞」，不倒裝；引述語氣時態常需後移",
    "rules": [
        "疑問詞保留，但語序改為直述句：Do you know where he lives?（非 where does he live）",
        "Yes/No 問句改用 whether 或 if 引導",
        "引述時常見動詞：say、tell、ask、wonder、know、find out",
    ],
    "examples": [
        ("Do you know <em>when Grandma is going to</em> visit us?",
         "你知道奶奶什麼時候要來看我們嗎？"),
        ("Jane wants to know <em>whether the weather will be</em> fine tomorrow.",
         "Jane 想知道明天天氣是否會好。"),
        ("He told me <em>that Olivia likes nice surprises</em>, not scary ones.",
         "他告訴我 Olivia 喜歡美好的驚喜，不喜歡嚇人的。"),
    ],
    "q_stem": "Can you tell me _____ the museum closes today?",
    "options": ["(A) when does", "(B) when", "(C) that when", "(D) what time does"],
    "answer": "(B) when",
    "explanation": "間接問句的疑問詞後面應使用直述句語序（主詞 + 動詞），因此保留 when 並接 the museum closes（不倒裝）。(A)(D) 保留了 does 的倒裝語序，(C) that when 結構錯誤，均不符合間接問句規則。",
  },
  {
    "num": 3, "tier": 1, "freq": "10/13",
    "zh": "過去完成式",
    "en": "Past Perfect Tense",
    "formula": "S + had + 過去分詞（p.p.）",
    "struct_note": "常與 by the time、before、after、when、already、never 連用，強調「過去的過去」",
    "rules": [
        "主詞無論人稱皆用 had + p.p.",
        "表示在某個過去時間點「之前」就已完成的動作",
        "否定：had not（hadn't）+ p.p.",
    ],
    "examples": [
        ("Then he noticed the snake he <em>had just seen</em> crossing the road a moment ago.",
         "然後他注意到就是剛才他見過穿越馬路的那條蛇。"),
        ("By the time we arrived at the theater, the show <em>had already started</em>.",
         "當我們到達劇院的時候，表演已經開始了。"),
        ("She <em>had never eaten</em> sushi before she visited Japan for the first time.",
         "她在第一次造訪日本之前從未吃過壽司。"),
    ],
    "q_stem": "When Tom got to the station, the train _____ already left.",
    "options": ["(A) has", "(B) had", "(C) was", "(D) is"],
    "answer": "(B) had",
    "explanation": "「火車離開」發生在「Tom抵達車站」之前，屬於過去的過去，須用過去完成式 had + p.p.（had left）。(A) has 是現在完成式，(C) was 是過去進行式的助動詞，(D) is 是現在式，均不符合時態邏輯。",
  },
  {
    "num": 4, "tier": 1, "freq": "10/13",
    "zh": "條件句（If 子句）",
    "en": "Conditional Sentences",
    "formula": "第一類（真實）：If + 現在式, will / can + 原形V\n第二類（假設）：If + 過去式（were）, would / could + 原形V\n第三類（過去假設）：If + had p.p., would have + p.p.",
    "struct_note": "注意：If 子句中不使用 will 或 would",
    "rules": [
        "If 子句放句首時加逗號；放句尾時不需要逗號",
        "第二類條件句：be 動詞一律用 were（所有人稱）",
        "unless = if...not：Unless it rains = If it doesn't rain",
    ],
    "examples": [
        ("<em>If you are taking your pet with you</em>, please buy a seat for it at half price.",
         "如果你要帶寵物同行，請以半價為牠購買一個座位。（第一類）"),
        ("<em>If I were a bird</em>, I could fly to any place I wanted.",
         "如果我是一隻鳥，我就可以飛到任何我想去的地方。（第二類）"),
        ("<em>If she had studied harder</em>, she would have passed the exam.",
         "如果她當時用功一點，她就會通過那場考試了。（第三類）"),
    ],
    "q_stem": "If it _____ tomorrow, we will cancel the outdoor activity.",
    "options": ["(A) will rain", "(B) rained", "(C) rains", "(D) had rained"],
    "answer": "(C) rains",
    "explanation": "主要子句為 will cancel（未來式），表示這是真實可能發生的情況（第一類條件句）。第一類條件句的 If 子句用現在式表未來，故選 (C) rains。(A) If 子句不能用 will，(B) 過去式用於第二類假設語氣，(D) 過去完成式用於第三類，均不正確。",
  },
  {
    "num": 5, "tier": 1, "freq": "9/13",
    "zh": "過去習慣（used to / would）",
    "en": "Used to / Habitual Past",
    "formula": "S + used to + 原形V（過去習慣或狀態，現在已不如此）\nS + would + 原形V（過去反覆的動作，不表狀態）",
    "struct_note": "易混淆：be used to + V-ing = 習慣於某事（現在仍如此）",
    "rules": [
        "used to 可表過去習慣動作或過去狀態；would 只能表過去反覆動作",
        "否定：didn't use to 或 used not to",
        "疑問：Did + S + use to + V?",
    ],
    "examples": [
        ("In my school days, I <em>used to listen</em> to English radio programs every day.",
         "在我求學時期，我每天都會收聽英語廣播節目。"),
        ("She <em>used to be</em> very shy, but now she loves talking to new people.",
         "她以前非常害羞，但現在她喜歡與新朋友交談。"),
        ("My grandfather <em>would tell</em> me stories every night before I fell asleep.",
         "每天晚上在我入睡之前，我祖父都會跟我說故事。"),
    ],
    "q_stem": "The town _____ a small fishing village before the tourists started coming.",
    "options": ["(A) would be", "(B) used to be", "(C) is used to be", "(D) was used to be"],
    "answer": "(B) used to be",
    "explanation": "句子描述小鎮過去的「狀態」（是個漁村），現在已改變，應用 used to be。(A) would be 只能表過去反覆動作，不能表過去狀態，故不適合。(C)(D) is used to / was used to 後接 V-ing，表示「習慣於某事（現在仍如此）」，語意不符。",
  },
  {
    "num": 6, "tier": 1, "freq": "9/13",
    "zh": "關係子句",
    "en": "Relative Clauses (who, which, that, where, when)",
    "formula": "先行詞 + 關係代名詞（who / which / that / where / when）+ 子句",
    "struct_note": "指人：who / that；指物：which / that；指地點：where；指時間：when",
    "rules": [
        "限定性關係子句（不加逗號）：The woman who lives next door is kind.",
        "非限定性關係子句（加逗號，不用 that）：My sister, who lives in Taipei, is a nurse.",
        "關係代名詞作受詞時可省略：The book (that) I bought is good.",
    ],
    "examples": [
        ("The woman <em>who lives next door</em> to Ken is a stranger to him.",
         "住在 Ken 隔壁的那位女性對他來說是個陌生人。"),
        ("This is the place <em>where I first met</em> my best friend ten years ago.",
         "這就是十年前我第一次遇見我最好朋友的地方。"),
        ("I still remember the day <em>when I got</em> my exam results and cried with joy.",
         "我仍然記得拿到考試成績、喜極而泣的那一天。"),
    ],
    "q_stem": "The tall man _____ you saw at the bookstore is my high school teacher.",
    "options": ["(A) who", "(B) which", "(C) where", "(D) what"],
    "answer": "(A) who",
    "explanation": "先行詞 The tall man 是人，關係子句中作受詞（you saw him），應用 who 或 that。(A) who 正確（也可用 that，但選項無 that）。(B) which 用於指物，(C) where 是關係副詞指地點，(D) what 不是關係代名詞，不可引導修飾名詞的子句。",
  },
  {
    "num": 7, "tier": 2, "freq": "9/13",
    "zh": "情態助動詞",
    "en": "Modal Verbs (can, could, may, might, must, should, would)",
    "formula": "S + 情態助動詞 + 原形動詞",
    "struct_note": "情態助動詞本身無人稱變化，後面一律接原形動詞，且不與另一個助動詞連用",
    "rules": [
        "can / could：能力、許可、可能（could 更委婉或表過去）",
        "may / might：許可、可能（might 可能性更低）",
        "must：強烈義務（must not = 禁止）；推測時 = 一定是",
        "should：建議、輕度義務",
    ],
    "examples": [
        ("You <em>should</em> take an umbrella with you in case it rains this afternoon.",
         "你應該帶把雨傘，以防今天下午下雨。（建議）"),
        ("She <em>must</em> be very tired after such a long trip without any rest.",
         "在這麼長途的旅程且毫無休息之後，她一定非常疲倦。（推測）"),
        ("You <em>might</em> find parking more difficult in this city than in others.",
         "你可能會發現在這座城市停車比在其他地方都要困難。（低度可能）"),
    ],
    "q_stem": "Students _____ return library books within two weeks, or they will have to pay a fine.",
    "options": ["(A) may", "(B) might", "(C) must", "(D) would"],
    "answer": "(C) must",
    "explanation": "句子後半段 or they will have to pay a fine（否則就要罰款）表示強制性的規定，應使用表「義務」的 must。(A) may（可能/可以）、(B) might（也許）表可能性，語氣不夠強，(D) would 表意願或假設，均不符合規定的語意。",
  },
  {
    "num": 8, "tier": 2, "freq": "8/13",
    "zh": "動名詞與不定詞",
    "en": "Gerunds and Infinitives",
    "formula": "動名詞：V-ing（當名詞使用）\n不定詞：to + 原形V",
    "struct_note": "It is + adj. + to V；enjoy / avoid / mind + V-ing；want / decide / plan + to V",
    "rules": [
        "enjoy, finish, avoid, mind, keep, suggest, give up + V-ing",
        "want, decide, plan, promise, refuse, agree, hope + to V",
        "stop + V-ing（停止做）vs. stop + to V（停下來去做）",
    ],
    "examples": [
        ("<em>Making goulash</em> together is a special tradition — we don't talk much while cooking.",
         "一起做燉肉是我們家的特別傳統——煮飯的時候我們話不多。"),
        ("I enjoy <em>reading</em> books right before going to bed—it helps me relax.",
         "我喜歡在睡覺之前看書——這有助於我放鬆。"),
        ("It is important <em>to eat</em> a balanced breakfast every morning.",
         "每天早上都吃均衡的早餐是很重要的。"),
    ],
    "q_stem": "Would you mind _____ the window? It's very hot in here.",
    "options": ["(A) open", "(B) to open", "(C) opening", "(D) opened"],
    "answer": "(C) opening",
    "explanation": "mind 是只能接動名詞（V-ing）的動詞，固定搭配為 Would you mind + V-ing?（你介意……嗎？）故選 (C) opening。(A) 原形 open 不可直接接在 mind 後，(B) to open 是不定詞，mind 不接 to V，(D) opened 是過去分詞，文法錯誤。",
  },
  {
    "num": 9, "tier": 2, "freq": "8/13",
    "zh": "被動語態",
    "en": "Passive Voice (be + past participle)",
    "formula": "現在被動：am / is / are + p.p.\n過去被動：was / were + p.p.\n未來被動：will be + p.p.\n完成被動：has / have been + p.p.",
    "struct_note": "說明執行者時用 by + 受詞格：The book was written by her.",
    "rules": [
        "被動語態將動作「接受者」移至主詞位置，強調動作本身",
        "使役動詞被動：be made to + V（注意加 to！）",
    ],
    "examples": [
        ("Popular singers like A-mei <em>are often seen</em> here with their friends.",
         "像阿妹這樣的知名歌手常常被看到在這裡與朋友相聚。"),
        ("The little girl <em>was hit</em> by a truck on her way home and was badly hurt.",
         "那個小女孩在回家途中被一輛卡車撞倒，受了重傷。"),
        ("The new community center <em>will be built</em> next to the park starting next year.",
         "新的社區中心將從明年開始在公園旁邊興建。"),
    ],
    "q_stem": "English _____ as an official language in more than 50 countries around the world.",
    "options": ["(A) speaks", "(B) is spoken", "(C) was speaking", "(D) has spoken"],
    "answer": "(B) is spoken",
    "explanation": "主詞 English 是語言，是被「說」的對象（動作接受者），應使用被動語態。句子描述現在普遍的事實，故用現在被動式 is spoken。(A) speaks 是主動式，(C) 是過去進行式主動，(D) 是現在完成式主動，均不正確。",
  },
  {
    "num": 10, "tier": 2, "freq": "7/13",
    "zh": "附加問句",
    "en": "Question Tags",
    "formula": "肯定主句 + 否定附加問句（助動詞 + not + 主詞代名詞）\n否定主句 + 肯定附加問句",
    "struct_note": "例外：I am → aren't I?；命令句 → will you? / shall we?",
    "rules": [
        "She is kind, isn't she?（肯定→否定）",
        "He doesn't like it, does he?（否定→肯定）",
        "主句無助動詞時，用 do / does / did 形成附加問句",
    ],
    "examples": [
        ("We are still good friends, <em>aren't we</em>?",
         "我們還是好朋友，不是嗎？"),
        ("She can speak three languages fluently, <em>can't she</em>?",
         "她能流利地說三種語言，不是嗎？"),
        ("You haven't seen my bag anywhere, <em>have you</em>?",
         "你沒有在哪裡看到我的包包，有嗎？"),
    ],
    "q_stem": "Tom didn't finish his homework last night, _____?",
    "options": ["(A) didn't he", "(B) did he", "(C) doesn't he", "(D) does he"],
    "answer": "(B) did he",
    "explanation": "主句為否定句（didn't finish），附加問句應用肯定形式；助動詞取自主句的 didn't，因此附加問句為 did he（不加 not）。(A) didn't he 是否定附加問句，不符規則；(C)(D) doesn't/does 時態不對（主句用過去式 didn't）。",
  },
  {
    "num": 11, "tier": 2, "freq": "7/13",
    "zh": "比較級與最高級",
    "en": "Comparative & Superlative",
    "formula": "比較級：adj./adv. + -er + than / more + adj./adv. + than\n最高級：the + adj./adv. + -est / the most + adj./adv.\n同等比較：as + 原級 + as\n雙重比較：The + 比較級, the + 比較級",
    "struct_note": "",
    "rules": [
        "good / well → better → best",
        "bad / badly → worse → worst",
        "many / much → more → most；little → less → least",
    ],
    "examples": [
        ("<em>The more</em> you practice, <em>the better</em> you will become.",
         "你練習得越多，你就會變得越好。"),
        ("The paint on the wall is not <em>as bright as</em> it was ten years ago.",
         "牆上的漆不像十年前那麼亮了。"),
        ("For Mike, the price is <em>the most important</em> thing when he shops for jeans.",
         "對 Mike 來說，買牛仔褲時價格是最重要的事。"),
    ],
    "q_stem": "The weather today is _____ than yesterday—perfect for a picnic!",
    "options": ["(A) more good", "(B) the best", "(C) much better", "(D) more better"],
    "answer": "(C) much better",
    "explanation": "good 的比較級是 better（不規則），且有 than 引導比較，應用比較級。much 可用來強調比較級（much better = 好多了）。(A) more good 錯誤，good 的比較級不用 more；(B) the best 是最高級，不用 than；(D) more better 雙重比較，文法錯誤。",
  },
  {
    "num": 12, "tier": 2, "freq": "7/13",
    "zh": "假設語氣／wish 句型",
    "en": "Subjunctive / Wish / Would Rather",
    "formula": "現在遺憾：I wish + S + 過去式（be 動詞一律用 were）\n過去遺憾：I wish + S + had + p.p.\n未來期望：I wish + S + would + V\n寧願：would rather + V + than + V",
    "struct_note": "",
    "rules": [
        "wish 後接假設語氣，表達與現實相反的遺憾或願望",
        "be 動詞在假設語氣中一律用 were（包括 I、he、she）",
        "as if / as though + 過去式：好像……（He acts as if he were the boss.）",
    ],
    "examples": [
        ("I <em>wish I were</em> able to go back in time and do things differently.",
         "我希望我能夠回到過去，把事情做得不一樣。（對現在的遺憾）"),
        ("If only <em>I had studied</em> harder for the exam—I really regret it now.",
         "如果當初我能更用功準備考試就好了——我現在真的很後悔。（對過去的遺憾）"),
        ("I <em>would rather</em> stay home and read <em>than</em> go out in terrible weather.",
         "我寧願待在家讀書，也不想在這麼糟糕的天氣出門。"),
    ],
    "q_stem": "Amy doesn't have enough time to travel. She wishes she _____ more free time.",
    "options": ["(A) has", "(B) had", "(C) have", "(D) will have"],
    "answer": "(B) had",
    "explanation": "wish 後面接的子句表達與現在事實相反的遺憾，須使用過去式。Amy 現在沒有足夠時間（現在事實），所以 wish 後用過去式 had 表達遺憾。(A) has 是現在式，(C) have 是原形，(D) will have 是未來式，均不符合 wish 假設語氣的規則。",
  },
  {
    "num": 13, "tier": 3, "freq": "6/13",
    "zh": "使役動詞",
    "en": "Causative Verbs (make / let / have / get)",
    "formula": "make / let / have + 受詞 + 原形V\nget + 受詞 + to V\nhave / get + 受詞 + p.p.（被動意思）",
    "struct_note": "被動式：be made to + V（注意補回 to！）",
    "rules": [
        "make：強迫（The movie made me cry.）",
        "let：允許（Let me try.）",
        "have：安排（Have him call me.）",
        "get：說服/設法（She got him to help.）",
    ],
    "examples": [
        ("The teacher <em>made the students clean</em> the classroom after the art project.",
         "老師叫學生在美術課結束後把教室打掃乾淨。"),
        ("Mom <em>let me stay</em> up late to watch the final game of the World Cup.",
         "媽媽讓我熬夜看世界盃的決賽。"),
        ("She <em>had her car fixed</em> by the mechanic before the long road trip.",
         "她在長途旅行之前請技師把她的車修好。"),
    ],
    "q_stem": "The coach _____ all the players run ten laps before practice began.",
    "options": ["(A) let", "(B) made", "(C) got", "(D) had"],
    "answer": "(B) made",
    "explanation": "空格後接「受詞 + 原形V（run）」，make 表示「強迫/讓某人做某事」，且句意有強制性（跑十圈是訓練要求）。(A) let（允許）語意不符；(C) got 後應接 to run（to V）；(D) had 後接 to run（to V），故均不正確。",
  },
  {
    "num": 14, "tier": 3, "freq": "6/13",
    "zh": "分詞片語",
    "en": "Participle Phrases (-ing / -ed phrases)",
    "formula": "現在分詞（V-ing）：表主動或進行\n過去分詞（p.p.）：表被動或完成\n分詞片語放句首：Tired of waiting, she left.",
    "struct_note": "注意：分詞的邏輯主詞必須與主句主詞一致，避免懸垂分詞",
    "rules": [
        "名詞修飾（後位）：the man holding the bag = the man who is holding the bag",
        "名詞修飾（後位）：the book written by him = the book that was written by him",
        "常見開頭：Located, Based, Given, Covered, Tired, Excited + 主句",
    ],
    "examples": [
        ("The man <em>holding the goods</em> carefully packed them one by one into small boxes.",
         "那個拿著商品的男子小心地將它們一件件裝進小盒子裡。"),
        ("<em>Sitting by the window</em>, he quietly watched the sunset fade over the mountains.",
         "坐在窗邊，他靜靜地看著夕陽在山後漸漸落下。"),
        ("The words <em>printed on the package</em> were far too small for elderly people to read.",
         "包裝上印的字對老年人來說太小，根本看不清。"),
    ],
    "q_stem": "_____ about the upcoming trip, Lisa packed all her bags the night before.",
    "options": ["(A) Exciting", "(B) Excited", "(C) She excited", "(D) To excite"],
    "answer": "(B) Excited",
    "explanation": "Lisa 是感受到興奮的人（被某事物激起興奮感），應使用過去分詞 Excited（表被動/感受），且 Excited 的邏輯主詞與主句主詞 Lisa 一致。(A) Exciting 表「令人興奮的」是主動意義，用來形容事物而非人的感受；(C) 結構錯誤；(D) To excite 是不定詞，不構成分詞片語。",
  },
  {
    "num": 15, "tier": 3, "freq": "6/13",
    "zh": "強調句型／分裂句",
    "en": "Emphatic Structures / Cleft Sentences",
    "formula": "It is / was + 【強調部分】 + that / who + 【其餘句子】\nWhat + S + V + is / was + 【強調部分】\nAll + S + V + is / was + 名詞（最小化強調）",
    "struct_note": "強調主詞（人）時用 who；強調其他成分用 that",
    "rules": [
        "It is Mary who called me.（強調主詞）",
        "It was yesterday that I met him.（強調時間）",
        "What I need is more time.（What 分裂句，強調述語）",
    ],
    "examples": [
        ("<em>It is</em> the friendship <em>that</em> truly matters in the end—not money or fame.",
         "到最後，真正重要的是友情——而非金錢或名聲。"),
        ("<em>It was</em> her mother <em>who</em> gave her the most important advice in her life.",
         "給了她人生中最重要建議的人，正是她的母親。"),
        ("<em>What</em> surprised everyone most <em>was</em> how quickly he recovered from the injury.",
         "讓所有人最驚訝的，是他從受傷中恢復的速度之快。"),
    ],
    "q_stem": "It was Amy _____ first suggested holding the party outdoors.",
    "options": ["(A) that", "(B) who", "(C) which", "(D) when"],
    "answer": "(B) who",
    "explanation": "分裂句 It was + 強調部分 + that/who + 其餘句子，強調的是主詞 Amy（人），因此用 who（也可用 that，但 who 更準確）。(B) who 正確。(C) which 指物，不指人；(D) when 是關係副詞，指時間；若選 (A) that 也可接受，但選項中 who 更精確，為最佳答案。",
  },
]

TIER_LABELS = {1: "Tier 1（必考）", 2: "Tier 2（重要）", 3: "Tier 3（進階）"}
TIER_BG_HEX = {1: "EDE9FE", 2: "D1FAE5", 3: "FEF3C7"}  # light purple/teal/amber

# ── Build Document ──────────────────────────────────────────────────────────
doc = Document()

# Page margins
section = doc.sections[0]
section.top_margin    = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# Default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = BLACK
style._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')

# ── COVER PAGE ──────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, '國中教育會考 英語科 教學講義', color=GRAY, size=10)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, '15大高頻文法考點', bold=True, color=DARK_PUR, size=26)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, '102–114年試題統計 × 結構說明 × 例句翻譯 × 練習題', color=GRAY, size=12)

doc.add_paragraph()

# Tier legend
for tier, label in TIER_LABELS.items():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    color = TIER_COLOR[tier]
    add_run(p, f'● {label}', bold=True, color=color, size=11)

doc.add_paragraph()

# TOC
p = doc.add_paragraph()
add_run(p, '目錄', bold=True, size=12, color=DARK_PUR)

for u in UNITS:
    p = doc.add_paragraph()
    add_run(p, f"{u['num']:02d}  ", bold=True, color=TIER_COLOR[u['tier']], size=11)
    add_run(p, f"{u['zh']}  ", bold=True, size=11)
    add_run(p, u['en'], color=GRAY, size=10)

p = doc.add_paragraph()
add_run(p, '本講義依據 102–114 年國中教育會考英語科歷屆試題文法統計編製', color=GRAY, size=9)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ── UNITS ───────────────────────────────────────────────────────────────────
for u in UNITS:
    tier  = u['tier']
    color = TIER_COLOR[tier]
    bg    = TIER_BG_HEX[tier]

    # ── Unit header (coloured table) ──
    tbl = doc.add_table(rows=1, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = 'Table Grid'

    # Set header background
    hex_map = {1: "4C1D95", 2: "065F46", 3: "78350F"}
    hdr_hex = hex_map[tier]

    num_cell  = tbl.cell(0, 0)
    text_cell = tbl.cell(0, 1)
    freq_cell = tbl.cell(0, 2)

    for cell in (num_cell, text_cell, freq_cell):
        set_cell_shading(cell, hdr_hex)

    # Column widths
    num_cell.width  = Cm(1.5)
    text_cell.width = Cm(11)
    freq_cell.width = Cm(2.5)

    p = num_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, str(u['num']), bold=True, color=RGBColor(0xFF,0xFF,0xFF), size=18)

    p = text_cell.paragraphs[0]
    add_run(p, u['zh'] + '\n', bold=True, color=RGBColor(0xFF,0xFF,0xFF), size=14)
    add_run(p, u['en'], color=RGBColor(0xDD,0xDD,0xDD), size=9)

    p = freq_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(p, u['freq'] + '\n', bold=True, color=RGBColor(0xFF,0xFF,0xFF), size=14)
    add_run(p, '年份出現', color=RGBColor(0xCC,0xCC,0xCC), size=8)

    doc.add_paragraph()  # spacer

    # ── Structure box ──
    p = doc.add_paragraph()
    add_run(p, '句型結構', bold=True, color=color, size=9)
    set_shading(p, bg)

    for line in u['formula'].split('\n'):
        p = doc.add_paragraph()
        add_run(p, line, bold=True, size=11)
        set_shading(p, bg)

    if u['struct_note']:
        p = doc.add_paragraph()
        # Bold key words in struct_note
        add_run(p, u['struct_note'], size=10, color=BLACK)
        set_shading(p, bg)

    doc.add_paragraph()

    # ── Core rules ──
    p = doc.add_paragraph()
    add_run(p, '核心規則', bold=True, color=GRAY, size=9)

    for rule in u['rules']:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Cm(0.5)
        add_run(p, rule, size=10)

    doc.add_paragraph()

    # ── Section label: 例句 ──
    p = doc.add_paragraph()
    add_run(p, '例句', bold=True, color=RGBColor(0xFF,0xFF,0xFF), size=9)
    set_shading(p, hex_map[tier])

    for i, (en_html, zh) in enumerate(u['examples'], 1):
        p = doc.add_paragraph()
        add_run(p, f'{i}.  ', bold=True, color=color, size=10)
        add_html_text(p, en_html, base_color=DARK_PUR, base_size=11, em_color=color)

        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Cm(0.7)
        add_run(p2, zh, color=GRAY, size=10)

    doc.add_paragraph()

    # ── Practice question ──
    p = doc.add_paragraph()
    add_run(p, '練習題', bold=True, color=color, size=9)
    set_shading(p, bg)

    # Question stem (replace _____ with underline spaces)
    stem = u['q_stem'].replace('_____', '___________')
    p = doc.add_paragraph()
    set_shading(p, bg)
    add_run(p, stem, bold=True, size=11)

    # Options in a 2x2 grid via table
    opt_tbl = doc.add_table(rows=2, cols=2)
    for idx, opt in enumerate(u['options']):
        cell = opt_tbl.cell(idx // 2, idx % 2)
        cp = cell.paragraphs[0]
        add_run(cp, opt, size=11)
        set_cell_shading(cell, bg)

    doc.add_paragraph()

    # Answer & explanation
    p = doc.add_paragraph()
    add_run(p, '答案：', bold=True, color=color, size=10)
    add_run(p, u['answer'], bold=True, size=10)
    add_border_bottom(p)

    p = doc.add_paragraph()
    add_run(p, u['explanation'], size=10, color=RGBColor(0x37,0x41,0x51))

    # Page break between units (but not after last)
    if u['num'] < 15:
        doc.add_page_break()

# ── Save ────────────────────────────────────────────────────────────────────
out = r"E:\ClaudeCode Projects\cap-eng\grammar_teaching.docx"
doc.save(out)
print(f"Saved → {out}")
