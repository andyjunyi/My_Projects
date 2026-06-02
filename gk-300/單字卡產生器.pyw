"""
單字卡產生器
============
用法：雙擊執行，或在命令列輸入 python 單字卡產生器.py

需求：
  pip install python-docx
  （Python 內建 tkinter，Windows 預設已含）

產生格式比照 supportive_python.py 參考檔：
  單字(深藍) + 音標(橘紅) + 詞類(灰) + 中譯
  ▶ 英文例句 → 中文翻譯(綠色)
  【常見搭配片語】區塊
  【同反義詞】區塊
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import csv
import json
import os
import re
import threading

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Twips
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.text import WD_LINE_SPACING
except ImportError:
    import subprocess, sys
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, RGBColor, Twips
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.text import WD_LINE_SPACING


SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_DIR = os.path.join(SCRIPT_DIR, "輸出docx")

# 自動轉換路徑：WSL 下把 E:\ 轉成 /mnt/e/
_DEFAULT_CSV = r"E:\My_Projects\gk-300\gk_350.csv"
if os.path.exists("/mnt/"):
    _DEFAULT_CSV = _DEFAULT_CSV.replace("E:\\", "/mnt/e/").replace("E:/", "/mnt/e/")
    _DEFAULT_CSV = _DEFAULT_CSV.replace("\\", "/")

# 預設 JSON 資料來源（同反義、搭配片語）
_DEFAULT_JSON = os.path.join(os.path.dirname(SCRIPT_DIR), "gk350_online", "website", "words_p1.json")


# ─── Constants (比照 supportive_python.py) ──────────────────────────────────

FONT_NAME = "Arial Unicode MS"

C_WORD  = RGBColor(0x21, 0x5E, 0x99)   # deep blue
C_PHON  = RGBColor(0xBF, 0x4E, 0x14)   # orange-red
C_POS   = RGBColor(0x66, 0x66, 0x66)   # grey
C_TRANS = RGBColor(0x4E, 0xA7, 0x2E)   # green
C_SYN   = RGBColor(0xFF, 0x00, 0x00)   # red
C_FOOT  = RGBColor(0x99, 0x99, 0x99)   # light grey (footer)


def _set_all_fonts(run, name):
    """Apply font name to all four font slots."""
    rpr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"),    name)
    rFonts.set(qn("w:eastAsia"), name)
    rFonts.set(qn("w:hAnsi"),    name)
    rFonts.set(qn("w:cs"),       name)
    existing = rpr.find(qn("w:rFonts"))
    if existing is not None:
        rpr.remove(existing)
    rpr.insert(0, rFonts)


def _add_run(para, text, size_half_pt, color=None, bold=False):
    """Add a run with Arial Unicode MS in all font slots."""
    if not text:
        return
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size_half_pt / 2)
    if color:
        run.font.color.rgb = color
    _set_all_fonts(run, FONT_NAME)
    return run


def _set_para_spacing(para, before=None, line=None, line_rule=None):
    """Set paragraph spacing."""
    pPr  = para._p.get_or_add_pPr()
    spEl = pPr.find(qn("w:spacing"))
    if spEl is None:
        spEl = OxmlElement("w:spacing")
        pPr.append(spEl)
    if before is not None:
        spEl.set(qn("w:before"), str(before))
    if line is not None:
        spEl.set(qn("w:line"), str(line))
    if line_rule is not None:
        spEl.set(qn("w:lineRule"), line_rule)


def _set_para_indent(para, left=None, hanging=None):
    """Set paragraph indent."""
    pPr   = para._p.get_or_add_pPr()
    indEl = pPr.find(qn("w:ind"))
    if indEl is None:
        indEl = OxmlElement("w:ind")
        pPr.append(indEl)
    if left is not None:
        indEl.set(qn("w:left"), str(left))
    if hanging is not None:
        indEl.set(qn("w:hanging"), str(hanging))


# ─── POS / definition helper ────────────────────────────────────────────────

_POS_PATTERN = re.compile(r"^\(([^)]+)\)\s*")  # e.g. "(v) " or "(adj) "


def parse_pos_and_def(definition_zh):
    """Parse '(v) 适应；改编' → ('(v)', '适应；改编')."""
    m = _POS_PATTERN.match(definition_zh)
    if m:
        pos_tag = f"({m.group(1)})"
        rest = definition_zh[m.end():]
        # 移除尾部的（常考搭配：...）註解
        rest = re.sub(r"（[^）]*常考搭配[^）]*）", "", rest).strip()
        rest = re.sub(r"\([^)]*常考搭配[^)]*\)", "", rest).strip()
        return pos_tag, rest
    return "", definition_zh


# ─── Docx builders (比照 supportive_python.py) ──────────────────────────────

def _new_para(doc):
    """Add a new paragraph and return it."""
    return doc.add_paragraph()


def p_title(doc, word, phonetic, pos, definition):
    """Title row: word + phonetic + separator + POS + definition."""
    para = _new_para(doc)
    _add_run(para, word,          size_half_pt=40, color=C_WORD, bold=True)
    if phonetic:
        _add_run(para, "  " + phonetic, size_half_pt=32, color=C_PHON)
    _add_run(para, "  ",          size_half_pt=28)
    _add_run(para, pos,           size_half_pt=26, color=C_POS,  bold=True)
    _add_run(para, "  " + definition, size_half_pt=28)
    return para


def p_en(doc, text):
    """English example sentence with ▶ prefix."""
    para = _new_para(doc)
    _set_para_spacing(para, before=120, line=180, line_rule="auto")
    _set_para_indent(para, left=540, hanging=340)
    _add_run(para, "▶ " + text, size_half_pt=28)
    return para


def p_zh(doc, text):
    """Chinese translation."""
    para = _new_para(doc)
    _set_para_spacing(para, line=180, line_rule="auto")
    _set_para_indent(para, left=454)
    _add_run(para, text, size_half_pt=24, color=C_TRANS)
    return para


def p_header(doc, text, size_half_pt):
    """Section header (22 for collocations, 28 for synonyms)."""
    para = _new_para(doc)
    _set_para_spacing(para, before=160)
    _add_run(para, text, size_half_pt=size_half_pt, color=C_POS, bold=True)
    return para


def p_colloc(doc, en_phrase, zh_meaning):
    """Collocation entry: ▶ EN → ZH."""
    para = _new_para(doc)
    _set_para_spacing(para, before=40)
    _set_para_indent(para, left=567, hanging=340)
    _add_run(para, "▶ " + en_phrase,    size_half_pt=28, color=C_PHON)
    _add_run(para, "  → " + zh_meaning, size_half_pt=28, color=C_POS)
    return para


def p_syn(doc, synonyms, antonyms=None):
    """Synonym / antonym row: (同) words (反) words."""
    para = _new_para(doc)
    _set_para_spacing(para, before=160)
    _set_para_indent(para, left=340)
    _add_run(para, "(同)",            size_half_pt=28, color=C_SYN, bold=True)
    _add_run(para, " " + synonyms,    size_half_pt=28, color=C_SYN)
    if antonyms:
        _add_run(para, "\u3000(反)",  size_half_pt=28, color=C_SYN, bold=True)
        _add_run(para, " " + antonyms, size_half_pt=28, color=C_SYN)
    return para


def p_footer(doc, text):
    """Source/credit footer."""
    para = _new_para(doc)
    _set_para_spacing(para, before=200)
    _add_run(para, text, size_half_pt=18, color=C_FOOT, bold=False)
    return para


# ─── JSON lookup for enriched data ──────────────────────────────────────────

def load_json_words(json_path):
    """Load words from JSON file and index by word (lowercase)."""
    if not os.path.exists(json_path):
        return {}
    try:
        with open(json_path, encoding="utf-8") as f:
            data = json.load(f)
        return {w.get("w", "").strip().lower(): w for w in data}
    except Exception:
        return {}


def enrich_from_json(word, json_index):
    """Look up extra data (col, syn, ant) from JSON index."""
    entry = json_index.get(word.lower())
    if not entry:
        return {}, {}, {}
    
    # Collocations
    cols_raw = entry.get("col") or []
    cols = []
    for c in cols_raw:
        cols.append((c.get("phrase", ""), c.get("zh", "")))
    
    # Synonyms
    syns_raw = entry.get("syn") or []
    syns = [s.get("w", "") + ((" " + s.get("m", "")) if s.get("m") else "") for s in syns_raw]
    
    # Antonyms
    ants_raw = entry.get("ant") or []
    ants = [a.get("w", "") + ((" " + a.get("m", "")) if a.get("m") else "") for a in ants_raw]
    
    return cols, syns, ants


# ─── Docx generation ────────────────────────────────────────────────────────

def set_page_margins(doc, top=567, right=1134, bottom=567, left=1134):
    """Set page to A4 with margins."""
    section = doc.sections[0]
    section.page_width  = 11906
    section.page_height = 16838
    section.top_margin    = Twips(top)
    section.bottom_margin = Twips(bottom)
    section.left_margin   = Twips(left)
    section.right_margin  = Twips(right)


def generate_one_docx(csv_row: dict, out_path: str, json_index: dict = None):
    """
    直接產生單字卡 docx（不使用範本）。
    格式完全比照 supportive_python.py 參考檔。
    """
    word       = csv_row.get("word", "").strip()
    kk         = csv_row.get("kk_phonetics", "").strip()
    def_zh     = csv_row.get("definition_zh", "").strip()
    example_en = csv_row.get("example_sentence", "").strip()
    example_zh = csv_row.get("sentence_zh", "").strip()

    # 解析 POS 與中譯
    pos, definition = parse_pos_and_def(def_zh)

    # 產自 JSON 的擴充資料
    cols, syns, ants = [], [], []
    if json_index is not None:
        cols, syns, ants = enrich_from_json(word, json_index)

    # ── 建立 docx ─────────────────────────────────────────────
    doc = Document()

    # 移除預設空段落
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    set_page_margins(doc)

    # P1 – 標題列
    p_title(doc, word, kk, pos, definition)

    # P2 – 英文例句（只取第一個例句）
    if example_en:
        p_en(doc, example_en)

    # P3 – 中文翻譯
    if example_zh:
        p_zh(doc, example_zh)

    # P4–5 – 【常見搭配片語】
    if cols:
        p_header(doc, "【常見搭配片語】", size_half_pt=22)
        for en_ph, zh_ph in cols:
            if en_ph:
                p_colloc(doc, en_ph, zh_ph)

    # P4–6 – 【同反義詞】
    if syns or ants:
        p_header(doc, "【同反義詞】", size_half_pt=28)
        p_syn(doc, ", ".join(syns), ", ".join(ants) if ants else None)

    # Footer – 來源
    source = csv_row.get("source", "").strip()
    if source:
        p_footer(doc, f"資料來源：{source}")

    doc.save(out_path)


# ═══════════════════════════════════════════════════════════════
#  CSV 工具
# ═══════════════════════════════════════════════════════════════

def load_csv(path: str):
    rows = []
    with open(path, encoding="utf-8-sig") as f:
        reader = csv.DictReader(f)
        for row in reader:
            rows.append({k.strip(): v.strip() for k, v in row.items()})
    return rows


# ═══════════════════════════════════════════════════════════════
#  GUI
# ═══════════════════════════════════════════════════════════════

class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("單字卡產生器")
        self.resizable(False, False)
        self.configure(padx=24, pady=20)

        self.csv_path = tk.StringVar()
        self.csv_path.set(_DEFAULT_CSV)
        self.json_path = tk.StringVar()
        self.json_path.set(_DEFAULT_JSON)
        self.out_dir  = tk.StringVar(value=OUTPUT_DIR)
        self.all_rows = []
        self.json_index = {}

        self._build_ui()
        self._check_json_source()
        # 自動載入預設 CSV
        self.after(100, self._auto_load_csv)

    def _build_ui(self):
        tk.Label(self, text="單字卡產生器", font=("Arial", 18, "bold")).grid(
            row=0, column=0, columnspan=3, pady=(0, 16), sticky="w")

        # CSV 選擇
        tk.Label(self, text="CSV 檔案：", anchor="w").grid(row=1, column=0, sticky="w", pady=4)
        tk.Entry(self, textvariable=self.csv_path, width=44, state="readonly").grid(
            row=1, column=1, sticky="ew", padx=(0, 8))
        tk.Button(self, text="瀏覽…", command=self._pick_csv).grid(row=1, column=2)

        # JSON 資料來源（可選，用來取得搭配片語/同反義）
        tk.Label(self, text="JSON 來源：", anchor="w").grid(row=2, column=0, sticky="w", pady=4)
        tk.Entry(self, textvariable=self.json_path, width=44, state="readonly").grid(
            row=2, column=1, sticky="ew", padx=(0, 8))
        tk.Button(self, text="瀏覽…", command=self._pick_json).grid(row=2, column=2)

        # 輸出資料夾
        tk.Label(self, text="輸出資料夾：", anchor="w").grid(row=3, column=0, sticky="w", pady=4)
        tk.Entry(self, textvariable=self.out_dir, width=44, state="readonly").grid(
            row=3, column=1, sticky="ew", padx=(0, 8))
        tk.Button(self, text="瀏覽…", command=self._pick_outdir).grid(row=3, column=2)

        ttk.Separator(self, orient="horizontal").grid(
            row=4, column=0, columnspan=3, sticky="ew", pady=12)

        # 單字輸入區
        tk.Label(self, text="輸入單字：", anchor="w").grid(
            row=5, column=0, sticky="nw", pady=(2, 0))

        input_frame = tk.Frame(self)
        input_frame.grid(row=5, column=1, columnspan=2, sticky="ew", pady=(0, 4))

        self.word_text = tk.Text(input_frame, width=44, height=8,
                                 font=("Arial", 12), relief="solid",
                                 bd=1, padx=8, pady=6)
        self.word_text.pack(side="left", fill="both", expand=True)
        sb = tk.Scrollbar(input_frame, command=self.word_text.yview)
        sb.pack(side="right", fill="y")
        self.word_text.config(yscrollcommand=sb.set)

        tk.Label(self, text="每行一個單字\n（大小寫不拘）",
                 font=("Arial", 10), fg="#888888", justify="left").grid(
            row=6, column=0, sticky="nw", pady=(0, 4))

        # 比對結果提示
        self.match_var = tk.StringVar(value="")
        tk.Label(self, textvariable=self.match_var, anchor="w",
                 fg="#1a5fb4", font=("Arial", 10)).grid(
            row=6, column=1, columnspan=2, sticky="w")

        self.word_text.bind("<KeyRelease>", lambda e: self._update_match())

        ttk.Separator(self, orient="horizontal").grid(
            row=7, column=0, columnspan=3, sticky="ew", pady=12)

        # 進度條
        self.progress = ttk.Progressbar(self, orient="horizontal",
                                        mode="determinate", length=500)
        self.progress.grid(row=8, column=0, columnspan=3, sticky="ew")

        self.status_var = tk.StringVar(value="請先載入 CSV 檔案。")
        tk.Label(self, textvariable=self.status_var, anchor="w",
                 fg="#555555", font=("Arial", 10)).grid(
            row=9, column=0, columnspan=3, sticky="w", pady=(4, 12))

        # 產生按鈕
        self.gen_btn = tk.Button(
            self, text="▶  產生 docx 檔案",
            command=self._start_generate,
            bg="#1a5fb4", fg="white",
            font=("Arial", 12, "bold"),
            padx=20, pady=8, relief="flat", cursor="hand2")
        self.gen_btn.grid(row=10, column=0, columnspan=3, pady=(0, 4))

    # ── 事件 ────────────────────────────────────────────────────

    def _check_json_source(self):
        path = self.json_path.get()
        if os.path.exists(path):
            self.json_index = load_json_words(path)
            if self.json_index:
                self.status_var.set(f"已載入搭配片語/同反義資料 ({len(self.json_index)} 字)")

    def _pick_csv(self):
        path = filedialog.askopenfilename(
            title="選擇 CSV 檔案",
            filetypes=[("CSV 檔案", "*.csv"), ("所有檔案", "*.*")])
        if not path:
            return
        self.csv_path.set(path)
        self._load_csv_file()

    def _pick_json(self):
        path = filedialog.askopenfilename(
            title="選擇 JSON 資料來源（含 col/syn/ant）",
            filetypes=[("JSON 檔案", "*.json")])
        if not path:
            return
        self.json_path.set(path)
        self.json_index = load_json_words(path)
        count = len(self.json_index)
        self.status_var.set(f"已載入 JSON ({count} 字)" + ("（含搭配片語/同反義）" if count else ""))

    def _auto_load_csv(self):
        if self.csv_path.get():
            self._load_csv_file()

    def _load_csv_file(self):
        path = self.csv_path.get()
        if not os.path.exists(path):
            self.status_var.set(f"找不到 CSV 檔案：{path}")
            return
        if not self.out_dir.get():
            self.out_dir.set(os.path.dirname(path))
        try:
            self.all_rows = load_csv(path)
            self.status_var.set(f"已載入 {len(self.all_rows)} 筆資料。")
            self._update_match()
        except Exception as e:
            messagebox.showerror("載入失敗", str(e))

    def _pick_outdir(self):
        path = filedialog.askdirectory(title="選擇輸出資料夾")
        if path:
            self.out_dir.set(path)

    def _get_input_words(self):
        raw = self.word_text.get("1.0", tk.END)
        return [w.strip() for w in raw.splitlines() if w.strip()]

    def _update_match(self):
        if not self.all_rows:
            return
        words = self._get_input_words()
        if not words:
            self.match_var.set("")
            return
        row_map   = {r.get("word", "").strip().lower(): r for r in self.all_rows}
        found     = [w for w in words if w.lower() in row_map]
        not_found = [w for w in words if w.lower() not in row_map]
        parts = [f"找到 {len(found)} 個"]
        if not_found:
            parts.append(f"找不到：{', '.join(not_found)}")
        # 顯示是否有搭配片語
        enrich_count = sum(1 for w in found if w.lower() in self.json_index)
        if enrich_count:
            parts.append(f"（{enrich_count} 個有搭配片語/同反義）")
        self.match_var.set("　".join(parts))

    def _start_generate(self):
        if not self.csv_path.get():
            messagebox.showwarning("提示", "請先選擇 CSV 檔案。")
            return
        if not self.out_dir.get():
            messagebox.showwarning("提示", "請先選擇輸出資料夾。")
            return
        words = self._get_input_words()
        if not words:
            messagebox.showwarning("提示", "請輸入至少一個單字。")
            return
        self.gen_btn.config(state="disabled")
        threading.Thread(target=self._generate, args=(words,), daemon=True).start()

    def _generate(self, words):
        out_dir = self.out_dir.get()
        os.makedirs(out_dir, exist_ok=True)
        row_map = {r.get("word", "").strip().lower(): r for r in self.all_rows}
        total   = len(words)
        success = 0
        errors  = []

        self.progress["maximum"] = total
        self.progress["value"]   = 0

        for i, word in enumerate(words, 1):
            self.status_var.set(f"產生中… {i}/{total}：{word}")
            row = row_map.get(word.lower())
            if row is None:
                errors.append(f"{word}（CSV 中找不到）")
                self.progress["value"] = i
                continue
            try:
                out_path = os.path.join(out_dir, f"{word}.docx")
                generate_one_docx(row, out_path, json_index=self.json_index)
                success += 1
            except Exception as e:
                errors.append(f"{word}（{e}）")
            self.progress["value"] = i

        if errors:
            msg = f"完成！成功 {success} 個，失敗 {len(errors)} 個：\n" + "\n".join(errors)
            self.after(0, lambda: messagebox.showwarning("部分失敗", msg))
        else:
            msg = f"全部完成！共產生 {success} 個 docx 檔案。\n\n儲存位置：\n{out_dir}"
            self.after(0, lambda: messagebox.showinfo("完成", msg))

        self.status_var.set(f"完成！已產生 {success} 個檔案。")
        self.gen_btn.config(state="normal")


# ═══════════════════════════════════════════════════════════════
#  入口
# ═══════════════════════════════════════════════════════════════

if __name__ == "__main__":
    app = App()
    app.mainloop()
